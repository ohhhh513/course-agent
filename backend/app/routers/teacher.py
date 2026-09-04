"""
教师端接口：/teacher/*  /analysis/*  /question/*
全部从真实 DB 动态聚合，不再依赖 teacher_class_data 静态快照
"""
import json
from datetime import datetime, date, timedelta
from collections import defaultdict
from fastapi import APIRouter, Depends, Query, Body, Form, UploadFile, File, HTTPException
from sqlalchemy.orm import Session
from sqlalchemy import func, desc
from pydantic import BaseModel
from typing import Optional, List

from ..database import get_db
from ..models.user import User, TeacherClass
from ..models.alert import Alert
from ..models.question import Question
from ..models.practice import PracticeSession, AnswerRecord
from ..models.graph import LearningPath, GraphNode
from ..models.course import Resource, ResourceProgress
from ..middleware.auth import get_current_user
from ..schemas.common import ok, fail, list_response
from ..utils import loads
from ..media_utils import (
    BASE_DIR, UPLOADS_DIR, COVERS_DIR, mp4_duration, pdf_pages, pptx_pages,
    guess_type, parse_chapter, parse_title, save_upload_file, generate_cover,
)
from ..routers.practice import _calc_question_stats

router = APIRouter(prefix="/api/v1/teacher", tags=["教师端"])


# ------- Pydantic 请求体 -------
class AlertReviewReq(BaseModel):
    action: str           # confirm / ignore / annotate
    note: Optional[str] = None


class SendMessageReq(BaseModel):
    userId: str
    content: str


# ------- 班级 ↔ 学生 关联辅助 -------
def _get_class_info(db: Session, class_id: str):
    """返回 (class_name, [student_user_ids]) 或 (None, [])"""
    tc = db.query(TeacherClass).filter(TeacherClass.class_id == class_id).first()
    if not tc:
        return None, []
    class_name = tc.class_name
    students = db.query(User.user_id, User.name, User.student_no, User.avatar_char, User.avatar_color).filter(
        User.role == "student", User.class_name == class_name
    ).all()
    return class_name, students


def _student_lp_stats(db: Session, user_ids: list) -> dict:
    """给一批 user_id 返回 {user_id: {total, done, mastery_avg, doing_count}}"""
    if not user_ids:
        return {}
    from sqlalchemy import case
    rows = db.query(
        LearningPath.user_id,
        func.count(LearningPath.kp_id),
        func.sum(case((LearningPath.status == "done", 1), else_=0)),
        func.avg(LearningPath.mastery),
        func.sum(case((LearningPath.status == "doing", 1), else_=0)),
    ).filter(
        LearningPath.user_id.in_(user_ids),
    ).group_by(LearningPath.user_id).all()
    return {
        r[0]: {"total": r[1] or 0, "done": r[2] or 0, "mastery": round(r[3] or 0, 1), "doing": r[4] or 0}
        for r in rows
    }


def _student_answer_stats(db: Session, user_ids: list) -> dict:
    """返回 {user_id: {total, correct}}"""
    if not user_ids:
        return {}
    rows = db.query(
        AnswerRecord.user_id,
        func.count(AnswerRecord.id),
        func.sum(AnswerRecord.is_correct),
    ).filter(AnswerRecord.user_id.in_(user_ids)).group_by(AnswerRecord.user_id).all()
    return {r[0]: {"total": r[1] or 0, "correct": r[2] or 0} for r in rows}


# ========== 教师名下班级列表（全动态，供前端下拉框使用） ==========
@router.get("/classes")
def teacher_classes(
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user),
):
    rows = db.query(TeacherClass).filter(TeacherClass.teacher_user_id == user.user_id).all()
    items = [{"classId": tc.class_id, "name": tc.class_name or tc.class_id} for tc in rows]
    return ok(items)


# ========== 教学驾驶舱（全动态） ==========
@router.get("/dashboard")
def teacher_dashboard(
    classId: str = Query("CL2301"),
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user),
):
    class_name, students = _get_class_info(db, classId)
    if not class_name:
        return fail("班级不存在", 404)

    student_ids = [s.user_id for s in students]
    lp_stats = _student_lp_stats(db, student_ids)
    ans_stats = _student_answer_stats(db, student_ids)

    # ==== classOverview ====
    n = len(student_ids)
    all_lp = db.query(LearningPath).filter(LearningPath.user_id.in_(student_ids)).all()
    total_kp = len(set(r.kp_id for r in all_lp))  # 总 kp 数（去重）
    # 平均完成率：每个学生「已完成知识点数 / 本人总知识点数」的平均值（全量口径，含未开始）
    per_student_completion = [
        (lp_stats[sid].get("done", 0) / lp_stats[sid].get("total", 0))
        for sid in student_ids if lp_stats.get(sid, {}).get("total", 0) > 0
    ]
    completion_rate = round(sum(per_student_completion) / len(per_student_completion) * 100, 1) if per_student_completion else 0
    # 平均学习完成率：全班所有「学生×知识点」格子的完成率均值（含未开始的 0%，不再只统计有进度的格子）
    mastery_rate = round(sum(r.mastery or 0 for r in all_lp) / len(all_lp), 1) if all_lp else 0
    # 已完成学生（所有 kp 都 done）
    fully_done = sum(1 for sid in student_ids if lp_stats.get(sid, {}).get("doing", 0) == 0 and lp_stats.get(sid, {}).get("total", 0) > 0)
    # 需关注学生（mastery < 60 的 kp 数 >= 3）
    need_attention = sum(
        1 for sid in student_ids
        if sum(1 for r in all_lp if r.user_id == sid and (r.mastery or 0) < 60) >= 3
    )
    # 预警聚合：red/yellow alert 数
    open_alerts_all = db.query(Alert).filter(
        Alert.class_id == classId, Alert.status != "closed",
    ).all()
    alert_student_count = len(set(a.user_id for a in open_alerts_all))
    alert_ratio = round(alert_student_count / n * 100, 1) if n else 0

    # 目标达成度 = 全班达到「掌握率 ≥ 60」的知识点格子数 / 总格子数（全量口径，与学生端能力目标达成度一致）
    reached = sum(1 for r in all_lp if (r.mastery or 0) >= 60)
    avg_goal_achieve = round(reached / len(all_lp) * 100, 1) if all_lp else 0

    # 今日活跃学生数 + 提交次数
    # 活跃 = 今日有 PracticeSession 提交 或 今日有 AnswerRecord 答题
    today = date.today()
    today_start = datetime.combine(today, datetime.min.time())

    # PracticeSession 提交
    today_sessions = db.query(PracticeSession).filter(
        PracticeSession.user_id.in_(student_ids),
        PracticeSession.finished_at >= today_start,
    ).all()
    active_from_session = set(p.user_id for p in today_sessions)
    submit_today = len(today_sessions)

    # AnswerRecord 答题（也算活跃 + 提交次数）
    today_answers = db.query(AnswerRecord).filter(
        AnswerRecord.user_id.in_(student_ids),
        AnswerRecord.created_at >= today_start,
    ).all()
    active_from_answer = set(a.user_id for a in today_answers)
    submit_today += len(today_answers)

    # 合并活跃学生
    active_user_ids = active_from_session | active_from_answer

    # kp 聚合（给 kpRanking 用）
    # 每个学生每个知识点的掌握率（无学习记录视为 0），分母统一为全班人数 n，
    # 避免「未达标人数 / 学习过人数」与班级总人数口径不一致导致「数量对不上」。
    kp_stu = defaultdict(dict)  # kp_id -> {user_id: (mastery, status)}
    for r in all_lp:
        kp_stu[r.kp_id][r.user_id] = (r.mastery or 0, r.status)

    kp_id_name = {n.id: n.name for n in db.query(GraphNode).filter(GraphNode.graph_type == "knowledge").all()}

    # kpRanking 前端期望: [{ name, mastery, weakCount, students }]
    kp_ranking = []
    for kp_id, stu_map in kp_stu.items():
        # 全班每人该知识点掌握率（未学习者记为 0，也即未达标）
        rows = [stu_map.get(sid, (0, "todo")) for sid in student_ids]
        masteries = [m for m, _ in rows]
        done_cnt = sum(1 for _, s in rows if s == "done")
        weak_cnt = sum(1 for m in masteries if m < 60)   # 未达标：全班掌握率 < 60（含未学习）
        avg_m = round(sum(masteries) / len(masteries), 1) if masteries else 0
        kp_ranking.append({
            "kpId": kp_id,
            "kpName": kp_id_name.get(kp_id, kp_id),
            "name": kp_id_name.get(kp_id, kp_id),          # 前端取 k.name
            "completionRate": round(done_cnt / n * 100, 1) if n else 0,
            "avgMastery": avg_m,
            "mastery": avg_m,                                # 前端取 k.mastery（全班均值）
            "weakCount": weak_cnt,                          # 前端取 k.weakCount（全班未达标人数）
            "students": n,                                  # 前端取 k.students（全班人数，与 donut 一致）
        })
    kp_ranking.sort(key=lambda x: x["mastery"])  # 低的在前面
    kp_ranking = kp_ranking[:5]                    # Top 5

    # ==== liveFeed：前端期望 [{ type, level, text, meta, time }] ====
    live = []
    if open_alerts_all:
        # 优先放最新的 alert 事件
        recent_alerts = sorted(open_alerts_all, key=lambda a: a.created_at or datetime.min, reverse=True)[:2]
        for a in recent_alerts:
            stu_name = next((s.name for s in students if s.user_id == a.user_id), a.user_id)
            kp_name = a.kp_name or kp_id_name.get(a.kp_id or "", "") or "未知"
            level = "danger" if a.level == "red" else "warn" if a.level == "yellow" else "ok"
            live.append({
                "type": "alert", "level": level,
                "text": f"新增预警：{stu_name} · {kp_name}",
                "meta": a.title or a.type or a.desc or "未处理",
                "time": (a.created_at or datetime.now()).strftime("%H:%M"),
            })
    # 再加最近的答题事件
    recent_ans = db.query(AnswerRecord).filter(
        AnswerRecord.user_id.in_(student_ids),
    ).order_by(desc(AnswerRecord.created_at)).limit(4).all()
    for ar in recent_ans:
        stu_name = next((s.name for s in students if s.user_id == ar.user_id), ar.user_id)
        kp_name = kp_id_name.get(ar.kp_id or "", "") or ar.kp_id or "未知"
        live.append({
            "type": "submit",
            "level": "ok" if ar.is_correct else "warn",
            "text": f"{stu_name} 提交了「{kp_name}」相关题目",
            "meta": "回答正确" if ar.is_correct else "回答错误",
            "time": (ar.created_at or datetime.now()).strftime("%H:%M"),
        })
    # 按 time 字符串降序取前 6
    live.sort(key=lambda x: x["time"], reverse=True)
    live = live[:6]

    # ==== todos：前端期望 [{ level, type, title, desc, action, target }] ====
    todos = []
    # 1) open alerts（最紧急）
    for a in sorted(open_alerts_all, key=lambda x: ({"red":0,"yellow":1}[x.level], x.created_at or datetime.min))[:4]:
        stu_name = next((s.name for s in students if s.user_id == a.user_id), a.user_id)
        kp_name = a.kp_name or kp_id_name.get(a.kp_id or "", "") or ""
        level = "danger" if a.level == "red" else "warn"
        todos.append({
            "id": f"TD_ALERT_{a.alert_id}",
            "level": level,
            "type": "alert",
            "title": a.title or f"「{kp_name}」{a.type or '学习预警'}",
            "desc": f"{stu_name} · {a.desc or kp_name or ''}",
            "action": "去处理",
            "target": "alerts",
        })
    # 2) 进度滞后的学生（completion < 50%）
    low_students = [sid for sid, lp in lp_stats.items() if lp["total"] and lp["done"] / lp["total"] < 0.5]
    low_students.sort(key=lambda sid: lp_stats[sid]["done"] / lp_stats[sid]["total"])
    for sid in low_students[:2]:
        stu = next((s for s in students if s.user_id == sid), None)
        if not stu: continue
        todos.append({
            "id": f"TD_LAG_{sid}",
            "level": "brand",
            "type": "homework",
            "title": f"{stu.name} 学习进度滞后",
            "desc": f"完成率 {round(lp_stats[sid]['done'] / lp_stats[sid]['total'] * 100, 1)}%，建议课后辅导",
            "action": "发消息",
            "target": "students",
            "userId": sid,
            "userName": stu.name,
        })

    return ok({
        "classOverview": {
            "classId": classId,
            "className": class_name,
            "studentCount": n,
            "totalKpCount": total_kp,
            "avgCompletionRate": completion_rate,
            "avgMasteryRate": mastery_rate,
            "avgGoalAchieve": avg_goal_achieve,         # 对齐前端
            "deltaMastery": None,                        # 暂无可比数据，前端 U.delta(null) 渲染为空
            "deltaCompletion": None,
            "deltaGoal": None,
            "fullyDoneStudents": fully_done,
            "needAttention": need_attention,
            "alertStudentCount": alert_student_count,     # 对齐前端
            "alertRatio": alert_ratio,                    # 对齐前端
            "activeToday": len(active_user_ids),          # 今日活跃学生数
            "submitToday": submit_today,                  # 今日提交次数
            "updatedAt": datetime.now().strftime("%Y-%m-%d %H:%M"),
        },
        "kpRanking": kp_ranking,
        "liveFeed": live,
        "todos": todos,
    })


# ========== 学情监测看板 · 学生列表（全动态） ==========
@router.get("/students")
def students(
    classId: str = Query("CL2301"),
    alertLevel: str = Query("all"),
    keyword: str = Query(None),
    sortBy: str = Query(None),
    page: int = Query(1),
    size: int = Query(20),
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user),
):
    class_name, students = _get_class_info(db, classId)
    if not class_name:
        return fail("班级不存在", 404)

    student_ids = [s.user_id for s in students]
    lp_stats = _student_lp_stats(db, student_ids)
    ans_stats = _student_answer_stats(db, student_ids)

    # 预警聚合：每个学生的最高级别 + 数量
    alert_rows = db.query(Alert.user_id, Alert.level, Alert.status).filter(
        Alert.class_id == classId, Alert.status != "closed",
    ).all()
    alert_by_user = defaultdict(lambda: {"red": 0, "yellow": 0, "top": "green"})
    for a in alert_rows:
        u = alert_by_user[a.user_id]
        u[a.level] = u.get(a.level, 0) + 1
        if a.level == "red": u["top"] = "red"
        elif a.level == "yellow" and u["top"] != "red": u["top"] = "yellow"

    # lastActive: 每个学生最近活跃时间（practice_sessions 或 answer_records 取 max）
    ps_rows = db.query(PracticeSession.user_id, PracticeSession.finished_at).filter(
        PracticeSession.user_id.in_(student_ids),
    ).all()
    ar_rows_all = db.query(AnswerRecord.user_id, AnswerRecord.created_at).filter(
        AnswerRecord.user_id.in_(student_ids),
    ).all()
    last_by_user = defaultdict(lambda: None)
    for uid, ts in ps_rows:
        if ts and (last_by_user[uid] is None or ts > last_by_user[uid]): last_by_user[uid] = ts
    for uid, ts in ar_rows_all:
        if ts and (last_by_user[uid] is None or ts > last_by_user[uid]): last_by_user[uid] = ts

    def _fmt_last(ts):
        if not ts: return "无活动"
        delta = (datetime.utcnow() - ts).total_seconds()
        if delta < 60: return "刚刚"
        if delta < 3600: return f"{int(delta // 60)} 分钟前"
        if delta < 86400: return f"{int(delta // 3600)} 小时前"
        if delta < 7 * 86400: return f"{int(delta // 86400)} 天前"
        return ts.strftime("%m-%d")

    items = []
    for s in students:
        lp = lp_stats.get(s.user_id, {})
        ans = ans_stats.get(s.user_id, {})
        completion = round(lp.get("done", 0) / lp.get("total", 0) * 100, 1) if lp.get("total", 0) else 0
        mastery = lp.get("mastery", 0)
        a = alert_by_user.get(s.user_id, {"top": "green", "red": 0, "yellow": 0})
        target = round(mastery + 15, 1) if mastery < 60 else round(mastery + 5, 1)
        items.append({
            "userId": s.user_id,
            "name": s.name,
            "no": s.student_no or "",
            "avatar": s.avatar_char or s.name[:1],
            "avatarColor": s.avatar_color or "indigo",
            "completion": completion,
            "mastery": mastery,
            "goal": min(target, 100),
            "alertLevel": a["top"],
            "alertCount": a["red"] + a["yellow"],
            "alertRed": a["red"],
            "alertYellow": a["yellow"],
            "lastActive": _fmt_last(last_by_user.get(s.user_id)),
            "answers": ans.get("total", 0),
            "correctRate": round(ans.get("correct", 0) / ans.get("total", 0) * 100, 1) if ans.get("total", 0) else 0,
        })

    # 过滤 + 排序 + 分页
    if alertLevel and alertLevel != "all":
        items = [s for s in items if s.get("alertLevel") == alertLevel]
    if keyword:
        k = keyword.lower()
        items = [s for s in items if k in (s.get("name", "") + s.get("no", "")).lower()]
    if sortBy == "mastery":
        items.sort(key=lambda s: s["mastery"])
    elif sortBy == "completion":
        items.sort(key=lambda s: s["completion"])
    elif sortBy == "alert":
        items.sort(key=lambda s: ({'danger':0,'warn':1,'ok':2}[s['alertLevel']]))
    total = len(items)
    start = (page - 1) * size
    return ok(list_response(items[start:start + size], total))


# ========== 学情监测看板 · 热力图（全动态，支持切换） ==========
@router.get("/heatmap")
def heatmap(
    classId: str = Query("CL2301"),
    dimension: str = Query("kp"),      # kp | week
    type: str = Query("completion"),   # completion | mastery
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user),
):
    class_name, students = _get_class_info(db, classId)
    if not class_name:
        return fail("班级不存在", 404)

    student_ids = [s.user_id for s in students]
    type = (type or "completion").lower()
    show_mastery = type == "mastery"

    # kpAxis：班级所有 kp（从 learning_paths 去重）
    lp_rows = db.query(LearningPath.kp_id, LearningPath.name, LearningPath.chapter).filter(
        LearningPath.user_id.in_(student_ids),
    ).distinct().all()
    kp_ids = []
    kp_names = []
    chapters = []
    seen = set()
    for r in lp_rows:
        if r.kp_id not in seen:
            seen.add(r.kp_id)
            kp_ids.append(r.kp_id)
            kp_names.append(r.name or r.kp_id)
            chapters.append(r.chapter or "")

    if show_mastery:
        # 知识点掌握热力图：基于 answer_records 的答题正确率
        from ..routers.graph import _mastery_from_records
        ans_rows = db.query(
            AnswerRecord.user_id, AnswerRecord.kp_id,
            func.count(AnswerRecord.id), func.sum(AnswerRecord.is_correct),
        ).filter(
            AnswerRecord.user_id.in_(student_ids),
            AnswerRecord.kp_id.in_(kp_ids),
        ).group_by(AnswerRecord.user_id, AnswerRecord.kp_id).all()
        score_dict = {(r[0], r[1]): (r[3] / r[2] * 100) if r[2] else 0 for r in ans_rows}
        # mask：有答题记录才显示；无记录 = 灰色
        started_mask = [
            [bool(score_dict.get((s.user_id, kp_id)) is not None) for kp_id in kp_ids]
            for s in students
        ]
    else:
        # 学习完成热力图：基于 learning_paths.mastery（资源学习进度）
        lp_rows2 = db.query(LearningPath.user_id, LearningPath.kp_id, LearningPath.mastery, LearningPath.status).filter(
            LearningPath.user_id.in_(student_ids),
            LearningPath.kp_id.in_(kp_ids),
        ).all()
        score_dict = {(r.user_id, r.kp_id): r.mastery or 0 for r in lp_rows2}
        # mask：status != 'todo' 表示已开始学习
        lp_started = {(r.user_id, r.kp_id): True for r in lp_rows2 if r.status != "todo"}
        started_mask = [
            [bool(lp_started.get((s.user_id, kp_id), False)) for kp_id in kp_ids]
            for s in students
        ]

    # data[row][col] = score 0-100；无有效数据的格子 = None（前端渲染为灰色）
    data = []
    for i, s in enumerate(students):
        row = []
        for j, kp_id in enumerate(kp_ids):
            if not started_mask[i][j]:
                row.append(None)
                continue
            key = (s.user_id, kp_id)
            score = score_dict.get(key, 0)
            row.append(round(score, 1))
        data.append(row)

    # kpAvg：每列平均 —— 只统计有效分数
    kp_avg = []
    for j in range(len(kp_ids)):
        col = [data[i][j] for i in range(len(students))
               if data[i][j] is not None and data[i][j] >= 0]
        kp_avg.append(round(sum(col) / len(col), 1) if col else None)

    # 每个 kp 有多少学生有数据
    kp_start_count = [
        sum(1 for i in range(len(students)) if started_mask[i][j])
        for j in range(len(kp_ids))
    ]

    # 共性薄弱识别：只展示平均（学习完成率 / 知识掌握率）< 80% 的知识点。
    # 升序排列后取全部低于 80% 的（最多 5 个），避免把已达标的高分知识点混入"薄弱"列表。
    avg_idx = [(j, kp_avg[j]) for j in range(len(kp_ids))
              if kp_avg[j] is not None and kp_avg[j] < 80]
    avg_idx.sort(key=lambda x: x[1])
    weakest_idx = [j for j, _ in avg_idx[:5]]
    weakest = [{"kpId": kp_ids[j], "kpName": kp_names[j], "name": kp_names[j],
                "avg": kp_avg[j] if kp_avg[j] is not None else 0,
                "avgMastery": kp_avg[j] if kp_avg[j] is not None else 0,
                "startedCount": kp_start_count[j]} for j in weakest_idx]

    return ok({
        "type": type,
        "label": "知识掌握率" if show_mastery else "学习完成率",
        "dimension": dimension,
        "kpAxis": kp_names,
        "studentAxis": [{"userId": s.user_id, "name": s.name, "no": s.student_no or ""} for s in students],
        "kpAxisIds": kp_ids,
        "kpAxisChapters": chapters,
        "data": data,
        "startedMask": started_mask,
        "kpAvg": kp_avg,
        "weakest": weakest,
    })


@router.get("/students/{user_id}/profile")
def student_profile(
    user_id: str,
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user),
):
    """个体学情详情（对齐 monitor.js openProfile 期望字段）"""
    stu = db.query(User).filter(User.user_id == user_id).first()
    if not stu:
        return fail("学生不存在", 404)

    lp_rows = db.query(LearningPath).filter(LearningPath.user_id == user_id).all()
    ans_rows = db.query(AnswerRecord).filter(AnswerRecord.user_id == user_id).all()

    total = len(lp_rows)
    done = [r for r in lp_rows if r.status == "done"]
    doing = [r for r in lp_rows if r.status == "doing"]
    completion = round(len(done) / total * 100, 1) if total else 0
    active = [r for r in lp_rows if (r.mastery or 0) > 0]
    mastery = round(sum(r.mastery or 0 for r in active) / len(active), 1) if active else 0
    accuracy = round(sum(1 for a in ans_rows if a.is_correct) / len(ans_rows) * 100, 1) if ans_rows else 0
    goal = round(mastery / 80 * 100, 1) if mastery else 0

    # 班级排名（按 mastery）
    class_name = stu.class_name or ""
    class_students = db.query(User).filter(User.role == "student", User.class_name == class_name).all()
    class_ids = [c.user_id for c in class_students]
    class_lps = db.query(LearningPath).filter(LearningPath.user_id.in_(class_ids)).all()
    class_masteries = defaultdict(list)
    for r in class_lps:
        if (r.mastery or 0) > 0:
            class_masteries[r.user_id].append(r.mastery or 0)
    class_avg = {}
    for uid, vals in class_masteries.items():
        class_avg[uid] = round(sum(vals) / len(vals), 1)
    sorted_class = sorted(class_avg.items(), key=lambda x: -x[1])
    rank = next((i + 1 for i, (uid, _) in enumerate(sorted_class) if uid == user_id), len(class_ids))

    # kp_detail：每个 kp 的掌握率 + 练习数 + 错题 + 时长
    ans_by_kp = defaultdict(list)
    for a in ans_rows: ans_by_kp[a.kp_id].append(a)
    kp_id_name = {n.id: n.name for n in db.query(GraphNode).filter(GraphNode.graph_type == "knowledge").all()}
    kp_detail = []
    for r in lp_rows:
        ar_list = ans_by_kp.get(r.kp_id, [])
        wrong = sum(1 for a in ar_list if not a.is_correct)
        minutes = sum(a.duration_seconds or 0 for a in ar_list) // 60
        m = round(r.mastery or 0, 1)
        level = "danger" if m < 60 else "warn" if m < 80 else "ok"
        kp_detail.append({
            "kpId": r.kp_id,
            "name": r.name or r.kp_id,
            "mastery": m, "level": level,
            "questions": len(ar_list), "wrong": wrong, "minutes": minutes,
        })
    kp_detail.sort(key=lambda x: x["mastery"])

    # wrongDetail：错题集中的 kp
    wrong_by_kp = defaultdict(list)
    for a in ans_rows:
        if not a.is_correct: wrong_by_kp[a.kp_id].append(a)
    wrong_detail = []
    for kp_id, wlist in sorted(wrong_by_kp.items(), key=lambda x: -len(x[1]))[:5]:
        wrong_detail.append({
            "kp": kp_id_name.get(kp_id, kp_id or "未知"),
            "qId": wlist[0].q_id or "-",
            "count": len(wlist),
            "errorType": wlist[0].error_type or "未分类",
        })

    # studyTimeDist：按真实答题时段（answer_records.created_at 的小时）聚合到 7 个时段桶；
    # 无记录则为全 0（真实空态，而非伪造分布）
    _buckets = [("0-2", 0, 2), ("2-6", 2, 6), ("6-9", 6, 9),
                ("9-12", 9, 12), ("12-14", 12, 14), ("14-18", 14, 18), ("18-22", 18, 22)]
    _hour_counts = [0] * 24
    for a in ans_rows:
        if a.created_at:
            _hour_counts[a.created_at.hour] += 1
    study_time_dist = {
        "xAxis": [b[0] for b in _buckets],
        "data": [sum(_hour_counts[b[1]:b[2]]) for b in _buckets],
    }

    # activityTrend：近 14 天
    today = date.today()
    act_x, act_min, act_q = [], [], []
    total_minutes = 0
    for offset in range(13, -1, -1):
        d = today - timedelta(days=offset)
        day_start = datetime.combine(d, datetime.min.time())
        day_end = datetime.combine(d, datetime.max.time())
        day_ps = [p for p in db.query(PracticeSession).filter(PracticeSession.user_id == user_id).all()
                  if p.finished_at and day_start <= p.finished_at <= day_end]
        mins = sum(p.duration_seconds or 0 for p in day_ps) // 60
        qs = sum(1 for a in ans_rows if a.created_at and day_start <= a.created_at <= day_end)
        act_x.append(d.strftime("%m-%d"))
        act_min.append(mins)
        act_q.append(0)  # AI 提问暂无真实来源
        total_minutes += mins

    return ok({
        "userId": user_id,
        "name": stu.name,
        "no": stu.student_no or "",
        "className": class_name,
        "metrics": {
            "completion": completion,
            "mastery": mastery,
            "goal": goal,
            "accuracy": accuracy,
            "rank": rank,
            "totalStudents": len(class_ids),
            "studyMinutes": total_minutes,
        },
        "studyTimeDist": study_time_dist,
        "activityTrend": {"xAxis": act_x, "minutes": act_min, "questions": act_q},
        "kpDetail": kp_detail,
        "wrongDetail": wrong_detail,
    })


# ========== 预警管理（已动态，无需改） ==========
@router.get("/alerts")
def teacher_alerts(
    classId: str = Query("CL2301"),
    level: str = Query("all"),
    status: str = Query("all"),
    type: str = Query(None),
    kpId: str = Query(None),
    keyword: str = Query(None),
    page: int = Query(1),
    size: int = Query(20),
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user),
):
    q = db.query(Alert)
    if classId and classId != "all":
        q = q.filter(Alert.class_id == classId)
    if level and level != "all":
        q = q.filter(Alert.level == level)
    if status and status != "all":
        q = q.filter(Alert.status == status)
    if kpId:
        q = q.filter(Alert.kp_id == kpId)
    rows = q.order_by(Alert.created_at.desc()).all()
    if keyword:
        k = keyword.lower()
        rows = [a for a in rows if k in ((a.kp_name or "") + (a.type or "") + (a.title or "")).lower()]

    # 学生 id → name 映射
    student_name_map = {u.user_id: u.name for u in db.query(User.user_id, User.name).filter(User.role == "student").all()}

    TYPE_LABEL = {
        "mastery_low": "掌握率偏低",
        "progress_lag": "学习进度滞后",
        "error_cluster": "错题集中爆发",
        "resolved": "预警已解除",
    }

    def alert_to_dict(a: Alert) -> dict:
        # 趋势数据：取该学生该 kp 最近 5 次 answer_records 的正确率
        trend_data = []
        ars = db.query(AnswerRecord).filter(
            AnswerRecord.user_id == a.user_id,
            AnswerRecord.kp_id == a.kp_id,
        ).order_by(AnswerRecord.created_at.desc()).limit(5).all()
        ars = list(reversed(ars))  # 时间升序
        if ars:
            # 用每 1 次正确率（单点）→ 填充到 5 个点
            for i in range(5):
                if i < len(ars):
                    trend_data.append(100 if ars[i].is_correct else 0)
                else:
                    trend_data.append(None)
        else:
            trend_data = [None] * 5

        return {
            "alertId": a.alert_id,
            "level": "red" if a.level == "red" else "yellow" if a.level == "yellow" else "green",
            "type": a.type or "",
            "typeLabel": TYPE_LABEL.get(a.type, a.type or "学习预警"),
            "title": a.title or "",
            "desc": a.desc or "",
            "trigger": a.trigger or "",
            "student": student_name_map.get(a.user_id, a.user_id),
            "kp": a.kp_name or "",
            "kpId": a.kp_id or "",
            "detail": loads(a.detail_json) or {},
            "trendData": trend_data,
            "createdAt": a.created_at.strftime("%Y-%m-%d %H:%M") if a.created_at else "",
            "status": "open" if a.status in ("read", "pending") else a.status,
            "note": a.note or "",
        }
    total = len(rows)
    start = (page - 1) * size
    return ok(list_response([alert_to_dict(a) for a in rows[start:start + size]], total))


@router.put("/alerts/{alert_id}/review")
def review_alert(
    alert_id: str,
    req: AlertReviewReq,
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user),
):
    alert = db.query(Alert).filter(Alert.alert_id == alert_id).first()
    if not alert:
        return fail("预警不存在", 404)
    # 越权校验：只能复核本班预警
    teacher_class_ids = {
        tc.class_id
        for tc in db.query(TeacherClass).filter(TeacherClass.teacher_user_id == user.user_id).all()
    }
    if alert.class_id not in teacher_class_ids:
        return fail("无权复核非本班预警", 403)
    alert.status = "ignored" if req.action == "ignore" else "reviewed"
    if req.note:
        alert.note = req.note
    alert.reviewed_at = datetime.utcnow()
    db.commit()
    return ok({"alertId": alert_id, "status": alert.status, "note": alert.note or ""})


@router.post("/messages")
def send_message(
    req: SendMessageReq,
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user),
):
    from ..models.alert import Message
    import uuid
    # 越权校验：只能向本班学生发送私信
    student = db.query(User).filter(User.user_id == req.userId, User.role == "student").first()
    if not student:
        return fail("接收学生不存在", 404)
    teacher_class_names = {
        tc.class_name
        for tc in db.query(TeacherClass).filter(TeacherClass.teacher_user_id == user.user_id).all()
    }
    if student.class_name not in teacher_class_names:
        return fail("无权向非本班学生发送私信", 403)
    msg = Message(
        msg_id="MSG" + uuid.uuid4().hex[:12],
        user_id=req.userId,
        from_user=user.user_id,
        from_name=user.name,
        title="教师私信",
        content=req.content,
    )
    db.add(msg)
    db.commit()
    return ok({"msgId": msg.msg_id, "to": req.userId, "sentAt": msg.created_at.strftime("%Y-%m-%dT%H:%M:%S")})


# ========== /analysis/* 错题归因（动态计算） ==========
analysis_router = APIRouter(prefix="/api/v1/analysis", tags=["归因分析"])


@analysis_router.get("/errors")
def analysis_errors(
    classId: str = Query("CL2301"),
    chapter: str = Query(None),
    kpId: str = Query(None),
    timeRange: str = Query(None),
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user),
):
    """错题归因分析 — 完整对齐 analysis.js 字段期望"""
    class_name, students = _get_class_info(db, classId)
    if not class_name:
        return fail("班级不存在", 404)
    student_ids = [s.user_id for s in students]
    kp_id_name = {n.id: n.name for n in db.query(GraphNode).filter(GraphNode.graph_type == "knowledge").all()}

    q = db.query(AnswerRecord).filter(AnswerRecord.user_id.in_(student_ids))
    if kpId: q = q.filter(AnswerRecord.kp_id == kpId)
    if chapter:
        lp_kps = [r[0] for r in db.query(LearningPath.kp_id).filter(
            LearningPath.user_id.in_(student_ids), LearningPath.chapter == chapter).distinct().all()]
        if lp_kps: q = q.filter(AnswerRecord.kp_id.in_(lp_kps))
    if timeRange == "7d": q = q.filter(AnswerRecord.created_at >= datetime.utcnow() - timedelta(days=7))
    elif timeRange == "30d": q = q.filter(AnswerRecord.created_at >= datetime.utcnow() - timedelta(days=30))

    rows = q.all()
    total = len(rows)
    wrong = [r for r in rows if not r.is_correct]

    # ----- topWrongQuestions -----
    q_cnt = defaultdict(int)
    q_total = defaultdict(int)
    for r in rows:
        q_total[r.q_id] += 1
        if not r.is_correct: q_cnt[r.q_id] += 1
    top_q_ids = sorted(q_cnt.items(), key=lambda x: x[1], reverse=True)[:5]
    top_q_rows = db.query(Question).filter(Question.q_id.in_([x[0] for x in top_q_ids])).all()
    q_map = {x.q_id: x for x in top_q_rows}
    top_wrong = []
    for qid, wc in top_q_ids:
        qobj = q_map.get(qid)
        if not qobj: continue
        t = q_total.get(qid, 1)
        top_wrong.append({
            "qId": qid,
            "stem": qobj.stem,
            "kp": kp_id_name.get(qobj.kp_id, qobj.kp_id or "未知"),
            "kpId": qobj.kp_id,
            "difficulty": qobj.difficulty,
            "mainWrongOption": "B" if wc > t * 0.6 else "C",  # 简化：无真实选项正确率时用启发式
            "wrongRate": round(wc / t * 100, 1),
            "count": wc,
        })

    # ----- weakChain（知识点薄弱链路 root→mid→leaf）-----
    # 取班级 mastery 最低的 3 个 kp 构成链路
    class_lps = db.query(LearningPath).filter(LearningPath.user_id.in_(student_ids)).all()
    kp_mastery = defaultdict(list)
    for r in class_lps:
        if (r.mastery or 0) > 0: kp_mastery[r.kp_id].append(r.mastery or 0)
    kp_avg_list = [(kid, round(sum(v) / len(v), 1)) for kid, v in kp_mastery.items()]
    kp_avg_list.sort(key=lambda x: x[1])
    chain = kp_avg_list[:3] if kp_avg_list else []
    chain_nodes = []
    for kid, m in chain:
        chain_nodes.append({
            "kpId": kid,
            "name": kp_id_name.get(kid, kid),
            "mastery": m,
        })
    while len(chain_nodes) < 3:
        chain_nodes.append({"kpId": "", "name": "—", "mastery": 0})

    # ----- causes（errorTypeDist → 前端成因）-----
    err_types = defaultdict(int)
    for r in wrong: err_types[r.error_type or "未分类"] += 1
    causes = []
    DESC_MAP = {
        "概念混淆": "对相近定义或边界条件理解模糊，易被干扰项诱导",
        "算法流程不清": "步骤记忆不准确，分支条件与循环边界常出错",
        "公式记忆错误": "关键公式系数或指数记错，导致计算偏差",
        "复杂度混淆": "时间/空间复杂度公式混用，best/worst case 判断不清",
        "指针操作不清": "指针移动顺序与目标结点删除/插入逻辑混乱",
        "定义记忆错误": "核心定义的适用场景或前提条件遗漏",
        "遍历顺序混淆": "前/中/后序、层序遍历的递归边界条件混淆",
        "WPL 计算失误": "带权路径长度累加顺序或权重理解偏差",
        "计算失误": "数值计算粗心导致",
        "未分类": "需进一步分析具体题目",
    }
    for et, cnt in sorted(err_types.items(), key=lambda x: -x[1])[:4]:
        pct = round(cnt / len(wrong) * 100, 1) if wrong else 0
        level = "danger" if pct >= 20 else "warn" if pct >= 10 else "ok"
        causes.append({
            "type": et,
            "title": et,
            "level": level,
            "pct": pct,
            "count": cnt,
            "desc": DESC_MAP.get(et, f"{et} 为主要错误成因之一，占错题 {pct}%"),
            "evidence": [f"该类型错题 {cnt} 道", f"占全部错题 {pct}%"],
            "advice": [
                f"针对「{et}」设计 3 道靶向补练题",
                "在 AI 答疑中嵌入相关概念辨析卡片",
                "课上组织 5 分钟错题归因讨论",
            ],
        })
    if not causes:
        causes = [{"type": "无错题", "title": "暂无错题成因", "level": "ok", "pct": 0, "count": 0,
                   "desc": "当前筛选范围内暂无错题记录", "evidence": ["班级答题全部正确"],
                   "advice": ["继续保持！"]}]

    # ----- commonVsIndividual -----
    # common：错题集中的 kp（只统计"学习过该 kp"的学生，即有 LP 记录且 status != 'todo'）
    kp_lp_started = defaultdict(set)  # 每个 kp 有多少学生真正学习过
    kp_wrong_students = defaultdict(set)
    for r in class_lps:
        if r.status != "todo":
            kp_lp_started[r.kp_id].add(r.user_id)
    for r in wrong:
        kp_wrong_students[r.kp_id].add(r.user_id)
    common = []
    for kid, uids in sorted(kp_wrong_students.items(), key=lambda x: -len(x[1])):
        affected = len(uids)
        started_cnt = len(kp_lp_started.get(kid, set()))
        if started_cnt == 0:
            continue  # 没有人学过这个 kp，不算共性薄弱
        ratio = round(affected / started_cnt * 100, 1)
        if ratio >= 30 and affected >= 2:
            common.append({
                "kpId": kid, "kp": kp_id_name.get(kid, kid or "未知"),
                "affected": affected, "startedCount": started_cnt, "ratio": ratio,
                "desc": f"{affected}/{started_cnt} 名学习过该知识点的学生出现错题，建议全班性巩固",
            })

    # individual：进度/掌握异常的个体学生（排除全部是 todo 的 kp）
    students_by_id = {s.user_id: s for s in students}
    lp_by_user = defaultdict(list)
    for r in class_lps: lp_by_user[r.user_id].append(r)
    individual = []
    for sid, lps in lp_by_user.items():
        # 只统计学习过（status != todo）的知识点
        active_lps = [r for r in lps if r.status != "todo"]
        if len(active_lps) < 3: continue  # 学习太少不判定
        masteries = [r.mastery or 0 for r in active_lps if (r.mastery or 0) > 0]
        avg_m = sum(masteries) / len(masteries) if masteries else 0
        done_cnt = sum(1 for r in active_lps if r.status == "done")
        total_active = len(active_lps)
        if avg_m < 40 or done_cnt / total_active < 0.3:
            stu = students_by_id.get(sid)
            issue = "掌握率极低" if avg_m < 40 else "进度严重滞后"
            individual.append({
                "userId": sid, "student": stu.name if stu else sid,
                "issue": issue,
                "desc": f"平均掌握率 {round(avg_m, 1)}% · 完成 {done_cnt}/{total_active} 学习点",
            })
    individual.sort(key=lambda x: -float(x["desc"].split("掌握率 ")[1].split("%")[0]) if "掌握率" in x["desc"] else 0)

    return ok({
        "scope": {"classId": classId, "chapter": chapter or "全部章节", "timeRange": timeRange or "全部"},
        "topWrongQuestions": top_wrong,
        "weakChain": {
            "root": chain_nodes[0],
            "mid": chain_nodes[1],
            "leaf": chain_nodes[2],
            "explain": f"班级「{class_name}」薄弱链路：{chain_nodes[0]['name']} → {chain_nodes[1]['name']} → {chain_nodes[2]['name']}，建议从前置知识点开始巩固",
        },
        "causes": causes,
        "commonVsIndividual": {"common": common, "individual": individual},
    })


@analysis_router.get("/weak-chain")
def analysis_weak_chain(
    classId: str = Query("CL2301"),
    kpId: str = Query(None),
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user),
):
    """弱链条分析 — 独立于 /errors，避免调用 __wrapped__"""
    class_name, students = _get_class_info(db, classId)
    if not class_name:
        return fail("班级不存在", 404)
    student_ids = [s.user_id for s in students]

    q = db.query(AnswerRecord).filter(AnswerRecord.user_id.in_(student_ids))
    if kpId:
        q = q.filter(AnswerRecord.kp_id == kpId)
    wrong = [r for r in q.all() if not r.is_correct]

    kp_wrong = defaultdict(int)
    for r in wrong:
        kp_wrong[r.kp_id] += 1
    if not kp_wrong:
        return ok(None)

    top_kp = sorted(kp_wrong.items(), key=lambda x: x[1], reverse=True)[0]
    kp_err = defaultdict(int)
    for r in wrong:
        if r.kp_id == top_kp[0]:
            kp_err[r.error_type or "未分类"] += 1
    top_err = sorted(kp_err.items(), key=lambda x: x[1], reverse=True)
    root = {"kpId": top_kp[0], "wrongCount": top_kp[1]}
    mid = {"errorType": top_err[0][0] if top_err else "未分类", "count": top_err[0][1] if top_err else 0}
    return ok({
        "root": root, "mid": mid,
        "leaf": {"count": len(wrong)},
        "explain": f"薄弱根因：{root['kpId']} ({root['wrongCount']} 错) → {mid['errorType']} ({mid['count']} 次)",
    })


@analysis_router.get("/causes")
def analysis_causes(
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user),
):
    """通用根因分类 — 保留枚举值，从 answer_records.error_type 统计 """
    rows = db.query(AnswerRecord.error_type, func.count(AnswerRecord.id)).group_by(AnswerRecord.error_type).all()
    total = sum(r[1] for r in rows)
    return ok({
        "categories": [
            {"name": r[0] or "未分类", "count": r[1], "pct": round(r[1] / total * 100, 1) if total else 0}
            for r in rows
        ],
    })


# ========== /question/* AI 出题（保留，仅 question/bank 为动态） ==========
question_router = APIRouter(prefix="/api/v1/question", tags=["AI 出题"])


class GenQuestionReq(BaseModel):
    materialIds: Optional[List[str]] = None
    kpIds: Optional[List[str]] = None
    types: Optional[List[str]] = None
    difficulty: int = 3
    count: int = 3
    skillId: Optional[str] = None
    requirement: Optional[str] = None


class ReviewReq(BaseModel):
    qIds: List[str]
    action: str


class UploadMaterialReq(BaseModel):
    name: str


class CreatePackReq(BaseModel):
    kpIds: List[str]
    userIds: Optional[List[str]] = None
    classId: Optional[str] = None
    count: int = 6
    difficulty: int = 3


@question_router.get("/gen/config")
def question_gen_config(
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user),
):
    """AI 出题配置：真实 kp 列表 + 题型选项"""
    kps = db.query(GraphNode).filter(GraphNode.graph_type == "knowledge").all()
    kp_options = sorted([{"kpId": k.id, "name": k.name, "chapter": k.chapter or ""} for k in kps],
                        key=lambda x: (x["chapter"], x["name"]))
    type_options = [
        {"key": "single", "name": "单选题"},
        {"key": "multi", "name": "多选题"},
        {"key": "judge", "name": "判断题"},
        {"key": "short", "name": "简答题"},
    ]
    return ok({"materials": [], "kpOptions": kp_options, "typeOptions": type_options})


@question_router.post("/materials")
def upload_material(
    req: UploadMaterialReq,
    user: User = Depends(get_current_user),
):
    """上传教材/课件 —— 生成占位素材条目，后续可接入真实文件解析"""
    import uuid
    file_id = "M" + uuid.uuid4().hex[:10]
    return ok({"fileId": file_id, "name": req.name, "status": "parsing", "progress": 0})


@question_router.post("/gen")
def ai_generate_questions(
    req: GenQuestionReq,
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user),
):
    """AI 生成习题 —— 从题库按 kp/type/difficulty 筛选，随机组合返回"""
    import uuid, random
    q = db.query(Question)
    if req.kpIds: q = q.filter(Question.kp_id.in_(req.kpIds))
    if req.types: q = q.filter(Question.type.in_(req.types))
    if req.difficulty: q = q.filter(Question.difficulty == req.difficulty)
    pool = q.filter(Question.status == "published").all()
    random.shuffle(pool)
    picked = pool[: req.count] if pool else []

    kp_id_name = {n.id: n.name for n in db.query(GraphNode).filter(GraphNode.graph_type == "knowledge").all()}

    questions = []
    for x in picked:
        questions.append({
            "qId": x.q_id, "type": x.type, "difficulty": x.difficulty,
            "stem": x.stem, "kpId": x.kp_id,
            "kpPath": loads(x.kp_path) or [kp_id_name.get(x.kp_id, x.kp_id or "未知")],
            "isKey": bool(x.is_key),
            "status": "pending",
            "options": loads(x.options) or [],
            "analysis": x.analysis or "",
            "classCorrectRate": round(_calc_question_stats(db, x.q_id)[0], 1),
        })
    return ok({
        "taskId": "GT" + uuid.uuid4().hex[:10],
        "count": len(questions),
        "questions": questions,
        "usedSkill": req.skillId or "SK004",
        "elapsedMs": 4200,
    })


@question_router.get("/bank")
def question_bank(
    kpId: str = Query(None),
    type: str = Query(None),
    status: str = Query("all"),
    difficulty: int = Query(None),
    keyword: str = Query(None),
    page: int = Query(1),
    size: int = Query(20),
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user),
):
    from ..models.question import Question
    q = db.query(Question)
    if kpId: q = q.filter(Question.kp_id == kpId)
    if type: q = q.filter(Question.type == type)
    if status and status != "all": q = q.filter(Question.status == status)
    if difficulty: q = q.filter(Question.difficulty == difficulty)
    if keyword:
        k = f"%{keyword}%"
        q = q.filter(Question.stem.like(k) | Question.q_id.like(k))
    all_items = q.all()

    # 批量算 correctRate（从 answer_records 聚合）
    q_ids = [x.q_id for x in all_items]
    ar_cnt = db.query(AnswerRecord.q_id, func.count(AnswerRecord.id), func.sum(AnswerRecord.is_correct)).filter(
        AnswerRecord.q_id.in_(q_ids),
    ).group_by(AnswerRecord.q_id).all()
    correct_rate_map = {}
    for qid, total, correct in ar_cnt:
        if total > 0:
            correct_rate_map[qid] = round(correct / total * 100, 1)

    kp_id_name = {n.id: n.name for n in db.query(GraphNode).filter(GraphNode.graph_type == "knowledge").all()}

    total = len(all_items)
    start = (page - 1) * size
    items = all_items[start:start + size]
    result = [
        {
            "qId": x.q_id, "type": x.type, "difficulty": x.difficulty,
            "status": x.status, "stem": x.stem,
            "kpId": x.kp_id,
            "kp": kp_id_name.get(x.kp_id, x.kp_id or "未知"),
            "correctRate": correct_rate_map.get(x.q_id),  # None = 暂无学生做过
            "isKey": bool(x.is_key),
            "options": loads(x.options) or [],
            "analysis": x.analysis,
            "kpPath": loads(x.kp_path) or [],
        }
        for x in items
    ]
    return ok(list_response(result, total))


@question_router.put("/{q_id}")
def update_question(
    q_id: str,
    body: dict = Body(...),
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user),
):
    q = db.query(Question).filter(Question.q_id == q_id).first()
    if not q:
        return fail("题目不存在", 404)

    # 基础字段直接赋值
    for key in ("stem", "answer", "analysis", "status", "difficulty", "kp_id", "is_key", "score", "type", "error_type"):
        if key in body:
            val = body[key]
            if key == "difficulty":
                val = int(val) if val is not None else q.difficulty
            elif key == "is_key":
                val = int(bool(val))
            elif key == "score":
                val = int(val)
            setattr(q, key, val)

    # options 字段需要序列化
    if "options" in body:
        opts = body["options"]
        q.options = json.dumps(opts, ensure_ascii=False) if isinstance(opts, list) else str(opts)

    # kp_path 字段需要序列化
    if "kp_path" in body:
        kpp = body["kp_path"]
        q.kp_path = json.dumps(kpp, ensure_ascii=False) if isinstance(kpp, list) else str(kpp)

    db.commit()
    return ok({"qId": q_id, "updated": True})


@question_router.post("/review")
def review_questions(
    req: ReviewReq,
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user),
):
    from ..models.question import Question
    q_ids = req.qIds
    if req.action in ("approve", "publish"):
        db.query(Question).filter(Question.q_id.in_(q_ids)).update({"status": req.action})
    elif req.action == "reject":
        db.query(Question).filter(Question.q_id.in_(q_ids)).update({"status": "archived"})
    db.commit()
    return ok({"affected": len(q_ids), "action": req.action})


@question_router.delete("/{q_id}")
def remove_question(
    q_id: str,
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user),
):
    from ..models.question import Question
    db.query(Question).filter(Question.q_id == q_id).delete()
    db.commit()
    return ok({"qId": q_id, "removed": True})


class ImportQuestionReq(BaseModel):
    questions: List[dict]


@question_router.post("/import")
def import_questions(
    req: ImportQuestionReq,
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user),
):
    """批量导入题目 —— 接收 JSON 数组，写入 questions 表"""
    import uuid
    success, failed = 0, 0
    errors = []
    kp_id_name = {n.id: n.name for n in db.query(GraphNode).filter(GraphNode.graph_type == "knowledge").all()}

    for i, q in enumerate(req.questions):
        try:
            stem = q.get("stem", "").strip()
            if not stem:
                failed += 1
                errors.append(f"第 {i+1} 题：题干为空")
                continue

            q_type = q.get("type", "single")
            difficulty = int(q.get("difficulty", 3))
            kp_id = q.get("kp_id", "")
            options = json.dumps(q.get("options", []), ensure_ascii=False) if isinstance(q.get("options"), list) else q.get("options", "[]")
            answer = q.get("answer", "")
            analysis = q.get("analysis", "")
            kp_path = json.dumps(q.get("kp_path", [kp_id_name.get(kp_id, kp_id)]), ensure_ascii=False)
            is_key = int(q.get("is_key", 0))

            new_q = Question(
                q_id="Q" + uuid.uuid4().hex[:8].upper(),
                course_id=q.get("course_id", "C2026DS001"),
                kp_id=kp_id,
                type=q_type,
                difficulty=difficulty,
                score=int(q.get("score", 5)),
                status=q.get("status", "pending"),
                stem=stem,
                options=options,
                answer=answer,
                analysis=analysis,
                kp_path=kp_path,
                pre_kp=json.dumps(q.get("pre_kp", []), ensure_ascii=False),
                post_kp=json.dumps(q.get("post_kp", []), ensure_ascii=False),
                is_key=is_key,
                class_correct_rate=0.0,
                avg_seconds=0,
                error_type=q.get("error_type", ""),
                source_ref_file="IMP",
                source_ref_locator="batch",
            )
            db.add(new_q)
            success += 1
        except Exception as e:
            failed += 1
            errors.append(f"第 {i+1} 题：{str(e)[:80]}")

    db.commit()
    return ok({
        "taskId": "IMP" + uuid.uuid4().hex[:8],
        "total": len(req.questions),
        "success": success,
        "failed": failed,
        "errors": errors[:10],
    })


# ====== 文件上传导入（Word / PDF / 图片）======
@question_router.post("/import/file")
async def import_questions_from_file(
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user),
):
    """从 Word(.docx) / PDF(.pdf) / 图片(.png/.jpg/.jpeg/.bmp) 文件解析题目并导入"""
    import os, tempfile, uuid as _u

    filename = file.filename or "upload"
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    # 写入临时文件
    tmp_dir = tempfile.gettempdir()
    tmp_path = os.path.join(tmp_dir, f"q_import_{_u.uuid4().hex[:12]}.{ext}")
    content = await file.read()
    with open(tmp_path, "wb") as f:
        f.write(content)

    try:
        # 提取文本
        if ext in ("docx", "doc"):
            text = _extract_text_docx(tmp_path)
        elif ext == "pdf":
            text = _extract_text_pdf(tmp_path)
        elif ext in ("png", "jpg", "jpeg", "bmp", "gif", "webp"):
            text = _extract_text_image(tmp_path)
        elif ext == "txt":
            with open(tmp_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
        else:
            os.remove(tmp_path)
            return fail(f"不支持的文件类型：{ext}。支持 .docx / .pdf / .txt / .png / .jpg / .jpeg")

        if not text or not text.strip():
            os.remove(tmp_path)
            return fail("文件内容为空或无法提取文本")

        # 从文本中解析结构化题目
        parsed = _parse_questions_from_text(text)
        if not parsed:
            os.remove(tmp_path)
            return ok({
                "taskId": "IMP" + _u.uuid4().hex[:8],
                "total": 0, "success": 0, "failed": 0,
                "warnings": ["未能从文件中自动识别出结构化题目。建议确保题目格式清晰（每题含题干、选项、答案）"],
                "rawPreview": text[:500],
                "hint": "系统识别格式如：1. 题干... A.选项 B.选项 C.选项 D.选项 答案:B",
            })

        # 写入 DB（复用 JSON import 的逻辑）
        success, failed = 0, 0
        errors = []
        kp_id_name = {n.id: n.name for n in db.query(GraphNode).filter(GraphNode.graph_type == "knowledge").all()}

        for i, q in enumerate(parsed):
            try:
                stem = q.get("stem", "").strip()
                if not stem:
                    failed += 1
                    errors.append(f"第 {i+1} 题：题干为空")
                    continue
                options_json = json.dumps(q.get("options", []), ensure_ascii=False) if isinstance(q.get("options"), list) else "[]"
                kp_id = q.get("kp_id", "")
                new_q = Question(
                    q_id="Q" + _u.uuid4().hex[:8].upper(),
                    course_id=q.get("course_id", "C2026DS001"),
                    kp_id=kp_id,
                    type=q.get("type", "single"),
                    difficulty=int(q.get("difficulty", 3)),
                    score=int(q.get("score", 5)),
                    status="pending",
                    stem=stem,
                    options=options_json,
                    answer=q.get("answer", ""),
                    analysis=q.get("analysis", ""),
                    kp_path=json.dumps(q.get("kp_path", [kp_id_name.get(kp_id, kp_id)]), ensure_ascii=False),
                    pre_kp=json.dumps(q.get("pre_kp", []), ensure_ascii=False),
                    post_kp=json.dumps(q.get("post_kp", []), ensure_ascii=False),
                    is_key=int(q.get("is_key", 0)),
                    class_correct_rate=0.0,
                    avg_seconds=0,
                    error_type=q.get("error_type", ""),
                    source_ref_file=filename,
                    source_ref_locator="file_import",
                )
                db.add(new_q)
                success += 1
            except Exception as e:
                failed += 1
                errors.append(f"第 {i+1} 题：{str(e)[:80]}")

        db.commit()
        os.remove(tmp_path)
        return ok({
            "taskId": "IMP" + _u.uuid4().hex[:8],
            "filename": filename, "format": ext,
            "total": len(parsed), "success": success, "failed": failed,
            "errors": errors[:10],
        })

    except Exception as e:
        # 清理临时文件
        if os.path.exists(tmp_path):
            try: os.remove(tmp_path)
            except: pass
        return fail(f"解析失败：{str(e)[:120]}")


def _extract_text_docx(path: str) -> str:
    """从 .docx 文件提取纯文本"""
    try:
        from docx import Document
        doc = Document(path)
        lines = []
        for para in doc.paragraphs:
            if para.text.strip():
                lines.append(para.text)
        # 表格里的文本也提取
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        lines.append(cell.text.strip())
        return "\n".join(lines)
    except ImportError:
        raise HTTPException(500, "python-docx 未安装，请先 pip install python-docx")


def _extract_text_pdf(path: str) -> str:
    """从 PDF 文件提取纯文本"""
    try:
        import pdfplumber
        lines = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                txt = page.extract_text() or ""
                if txt.strip():
                    lines.append(txt)
                # 表格文本
                for table in page.extract_tables() or []:
                    for row in table:
                        for cell in row:
                            if cell and str(cell).strip():
                                lines.append(str(cell).strip())
        return "\n".join(lines)
    except ImportError:
        raise HTTPException(500, "pdfplumber 未安装，请先 pip install pdfplumber")


def _extract_text_image(path: str) -> str:
    """从图片 OCR 提取文本（需要 Tesseract）"""
    try:
        import pytesseract
        from PIL import Image
        img = Image.open(path)
        # 自动检测中文 + 英文
        try:
            text = pytesseract.image_to_string(img, lang="chi_sim+eng")
        except pytesseract.pytesseract.TesseractNotFoundError:
            raise HTTPException(500,
                "Tesseract OCR 未安装。请：\n"
                "1. 下载安装：https://github.com/UB-Mannheim/tesseract/wiki\n"
                "2. 或运行：winget install UB-Mannheim.TesseractOCR\n"
                "3. 设置环境变量 TESSERACT_CMD 指向 tesseract.exe 路径\n"
                "4. 安装中文语言包 chi_sim")
        return text
    except ImportError:
        raise HTTPException(500, "pytesseract 或 Pillow 未安装，请先 pip install pytesseract pillow")


def _parse_questions_from_text(text: str) -> list:
    """从纯文本中解析结构化题目
    支持格式：
      1. 题干文本...
         A. 选项A
         B. 选项B
         C. 选项C
         D. 选项D
         答案: B
         解析: ...
      或：
      1) 题干... A) xxx B) xxx C) xxx D) xxx 答案：B
    """
    import re

    # 统一换行符
    text = text.replace("\r\n", "\n").replace("\r", "\n")

    # 按题号切分题目块
    # 匹配模式：1.  1)  1、  第1题
    q_blocks = re.split(r'\n(?=(?:第\s*)?\d+\s*(?:[.、\)）．]))', text)

    questions = []
    for block in q_blocks:
        block = block.strip()
        if not block or len(block) < 10:
            continue

        # 去除题号前缀
        block = re.sub(r'^(?:第\s*)?\d+\s*[.、\)）．]\s*', '', block)

        lines = [l.strip() for l in block.split("\n") if l.strip()]
        if not lines:
            continue

        stem_lines = []
        options = []
        answer = ""
        analysis = ""

        in_options = False
        current_opt_key = None
        current_opt_text = ""

        # 选项匹配: A.  A)  A、  A： 或 (A)
        opt_pattern = re.compile(r'^\(?([A-D])\)?[\.、\)）:：\s]')

        for line in lines:
            # 先检查是否是答案行
            ans_match = re.match(r'^(?:答案|参考答案|answer)[:：\s]*([A-Da-d]{1,4})', line, re.IGNORECASE)
            if ans_match:
                answer = ans_match.group(1).upper()
                continue

            # 检查是否是解析行
            ana_match = re.match(r'^(?:解析|分析|analysis|详解)[:：]\s*(.*)', line, re.IGNORECASE)
            if ana_match:
                analysis = ana_match.group(1)
                continue

            # 检查是否是新选项
            opt_match = opt_pattern.match(line)
            if opt_match:
                # 保存上一个选项
                if current_opt_key:
                    options.append({"key": current_opt_key, "text": current_opt_text.strip(),
                                    "right": current_opt_key and current_opt_key in answer})
                current_opt_key = opt_match.group(1).upper()
                current_opt_text = opt_pattern.sub('', line, count=1)
                in_options = True
            elif in_options and current_opt_key:
                # 选项换行
                current_opt_text += " " + line
            else:
                stem_lines.append(line)

        # 保存最后一个选项
        if current_opt_key:
            options.append({"key": current_opt_key, "text": current_opt_text.strip(),
                            "right": current_opt_key and current_opt_key in answer})

        stem = " ".join(stem_lines).strip()
        if len(stem) < 5:
            continue

        q_type = "single"
        if len(answer) > 1:
            q_type = "multiple"
        elif not options:
            q_type = "blank"

        questions.append({
            "stem": stem,
            "type": q_type,
            "difficulty": 3,
            "options": options,
            "answer": answer,
            "analysis": analysis,
        })

    return questions


@question_router.post("/packs")
def create_pack(
    req: CreatePackReq,
    user: User = Depends(get_current_user),
):
    """生成靶向补练包"""
    import uuid
    return ok({
        "packId": "PK" + uuid.uuid4().hex[:10],
        "kpIds": req.kpIds, "count": req.count, "pushed": True,
    })


# =============================================================================
# 教师资源管理：上传 / 删除 / 列表
# =============================================================================

def _next_res_id(db: Session) -> str:
    last = db.query(Resource.res_id).filter(Resource.res_id.like("R%")).order_by(Resource.res_id.desc()).first()
    n = int(last[0][1:]) + 1 if last else 1
    while db.query(Resource.res_id).filter(Resource.res_id == f"R{n:03d}").first():
        n += 1
    return f"R{n:03d}"


@router.get("/resources")
def teacher_resources(
    keyword: str = Query(None),
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user),
):
    """教师端：列出所有学习资源（含封面、观看次数、分类）"""
    q = db.query(Resource).order_by(Resource.res_id.asc())
    if keyword:
        like = f"%{keyword}%"
        q = q.filter(Resource.title.like(like) | Resource.kp.like(like) | Resource.source.like(like))
    rows = q.all()
    items = [
        {
            "resId": r.res_id, "type": r.type, "title": r.title,
            "kp": r.kp, "category": r.category,
            "duration": r.duration, "pages": r.pages, "count": r.count,
            "source": r.source, "views": r.views, "url": r.url or "",
            "cover": f"/assets/resources/covers/{r.res_id}.jpg",
        }
        for r in rows
    ]
    return ok(list_response(items, len(items)))


@router.post("/resources/upload")
async def upload_resource(
    title: str = Form(""),
    kp: str = Form(""),
    kp_id: str = Form(""),
    category: str = Form("other"),
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user),
):
    """教师上传资源：保存文件、解析时长/页数、生成封面、写入 Resource 表"""
    filename = file.filename or "upload.bin"
    rtype, ext = guess_type(filename)
    if not rtype:
        return fail("仅支持 MP4/PDF/PPT/DOC 文件", 400)

    res_id = _next_res_id(db)
    dest = save_upload_file(file.file, filename, res_id)

    # 解析元数据
    duration = ""
    pages = 0
    if rtype == "video":
        duration = mp4_duration(dest)
    elif rtype == "ppt":
        pages = pptx_pages(dest)
    elif rtype == "doc":
        pages = pdf_pages(dest) if ext == ".pdf" else 0

    title = title.strip() or parse_title(filename)
    kp = kp.strip() or parse_chapter(filename)

    url = f"/assets/resources/uploads/{res_id}/{filename}"
    r = Resource(
        res_id=res_id,
        course_id="C2026DS001",
        title=title,
        type=rtype,
        kp=kp,
        kp_id=kp_id or "",
        category=category or "other",
        duration=duration,
        pages=pages,
        count=0,
        source="教师上传",
        views=0,
        url=url,
    )
    db.add(r)
    db.commit()

    cover_url = generate_cover(res_id, title, rtype)

    return ok({
        "resId": res_id, "title": title, "type": rtype,
        "duration": duration, "pages": pages, "url": url,
        "cover": cover_url, "kp": kp, "category": category,
    })


@router.delete("/resources/{res_id}")
def delete_resource(
    res_id: str,
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user),
):
    """教师删除资源：删除 DB 记录、磁盘文件与封面"""
    r = db.query(Resource).filter(Resource.res_id == res_id).first()
    if not r:
        return fail("资源不存在", 404)

    # 删除文件
    try:
        if r.url:
            p = BASE_DIR / r.url.lstrip("/")
            if p.exists():
                p.unlink()
            d = p.parent
            if d.exists() and not any(d.iterdir()):
                d.rmdir()
    except Exception:
        pass

    # 删除封面
    try:
        c = COVERS_DIR / f"{res_id}.jpg"
        if c.exists():
            c.unlink()
    except Exception:
        pass

    # 级联删除学生的学习进度记录，避免外键约束失败
    try:
        db.query(ResourceProgress).filter(ResourceProgress.res_id == res_id).delete(synchronize_session=False)
    except Exception:
        pass

    db.delete(r)
    db.commit()
    return ok({"resId": res_id})


@router.get("/resources/kps")
def teacher_resource_kps(
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user),
):
    """返回所有知识点节点（用于上传时选择挂载知识点）"""
    nodes = db.query(GraphNode).filter(GraphNode.graph_type == "knowledge").order_by(
        GraphNode.chapter.asc(), GraphNode.id.asc()
    ).all()
    groups = {}
    for n in nodes:
        groups.setdefault(n.chapter or "其他", []).append({"kpId": n.id, "name": n.name})
    return ok({
        "chapters": [{"chapter": ch, "items": items} for ch, items in groups.items()]
    })
